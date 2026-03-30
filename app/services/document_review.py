from __future__ import annotations

import io
import mimetypes
import re
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from urllib.request import url2pathname

import httpx
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app.config import settings
from app.schemas import Attachment, CrossCheckResult, DocumentReviewResult, EvidenceRef, ReviewRecommendation, ReviewedDocument
from app.services.platform_registry import PlatformPackage

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None

try:  # pragma: no cover
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _normalize_entity(value: Any) -> str:
    if value is None:
        return ""
    cleaned = re.sub(r"[^A-Za-z0-9]+", " ", str(value)).strip().lower()
    tokens = [token for token in cleaned.split() if token]
    if not tokens:
        return ""
    return " ".join(sorted(tokens))


def _coerce_attachment_path(url: str) -> Path | None:
    if url.startswith("file://"):
        parsed = urlparse(url)
        return Path(url2pathname(parsed.path))
    path = Path(url)
    return path if path.exists() else None


def _read_attachment_bytes(attachment: Attachment) -> bytes:
    local = _coerce_attachment_path(attachment.url)
    if local is not None:
        return local.read_bytes()
    response = httpx.get(attachment.url, timeout=settings.request_timeout_s)
    response.raise_for_status()
    return response.content


def _attach_file_name(file_name: str, refs: list[EvidenceRef]) -> list[EvidenceRef]:
    for ref in refs:
        ref.file_name = file_name
    return refs


def _extract_pdf(data: bytes, file_name: str) -> tuple[str, list[EvidenceRef]]:
    reader = PdfReader(io.BytesIO(data))
    texts: list[str] = []
    evidence: list[EvidenceRef] = []
    for index, page in enumerate(reader.pages, start=1):
        text = _normalize_whitespace(page.extract_text() or "")
        if text:
            texts.append(text)
            evidence.append(EvidenceRef(file_name=file_name, page=index, snippet=text[:220]))
    return "\n".join(texts), evidence


def _extract_docx(data: bytes, file_name: str) -> tuple[str, list[EvidenceRef]]:
    document = Document(io.BytesIO(data))
    texts: list[str] = []
    evidence: list[EvidenceRef] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = _normalize_whitespace(paragraph.text)
        if text:
            texts.append(text)
            evidence.append(EvidenceRef(file_name=file_name, locator=f"paragraph:{index}", snippet=text[:220]))
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            row_text = " | ".join(_normalize_whitespace(cell.text) for cell in row.cells if cell.text)
            if row_text:
                rows.append(row_text)
        text = _normalize_whitespace(" ".join(rows))
        if text:
            texts.append(text)
            evidence.append(EvidenceRef(file_name=file_name, locator=f"table:{table_index}", snippet=text[:220]))
    return "\n".join(texts), evidence


def _extract_xlsx(data: bytes, file_name: str) -> tuple[str, list[EvidenceRef]]:
    workbook = load_workbook(io.BytesIO(data), data_only=True)
    texts: list[str] = []
    evidence: list[EvidenceRef] = []
    for sheet in workbook.worksheets:
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if not values:
                continue
            text = " | ".join(values)
            texts.append(text)
            evidence.append(EvidenceRef(file_name=file_name, locator=f"{sheet.title}:{row_index}", snippet=text[:220]))
    return "\n".join(texts), evidence


def _load_sidecar_ocr(attachment: Attachment) -> str:
    local = _coerce_attachment_path(attachment.url)
    if local is None:
        return ""
    sidecar = Path(f"{local}.ocr.txt")
    if not sidecar.exists():
        return ""
    return _normalize_whitespace(sidecar.read_text(encoding="utf-8"))


def _extract_image(data: bytes, attachment: Attachment) -> tuple[str, list[EvidenceRef]]:
    file_hint = _normalize_whitespace(Path(attachment.name).stem.replace("_", " ").replace("-", " "))
    texts: list[str] = [file_hint] if file_hint else []
    if Image is not None and pytesseract is not None:
        image = Image.open(io.BytesIO(data))
        ocr_text = _normalize_whitespace(pytesseract.image_to_string(image, lang="eng+chi_sim"))
        if ocr_text:
            texts.append(ocr_text)
    sidecar_text = _load_sidecar_ocr(attachment)
    if sidecar_text:
        texts.append(sidecar_text)
    snippet = " ".join(texts)[:220] if texts else attachment.name
    refs = [EvidenceRef(file_name=attachment.name, locator="image", snippet=snippet)]
    return "\n".join(texts), refs


def _extract_text(attachment: Attachment, data: bytes) -> tuple[str, list[EvidenceRef]]:
    mime = attachment.mime_type or mimetypes.guess_type(attachment.name)[0] or ""
    suffix = Path(attachment.name).suffix.lower()
    if mime == "application/pdf" or suffix == ".pdf":
        return _extract_pdf(data, attachment.name)
    if suffix == ".docx":
        return _extract_docx(data, attachment.name)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(data, attachment.name)
    if mime.startswith("image/") or suffix in {".png", ".jpg", ".jpeg"}:
        return _extract_image(data, attachment)
    text = _normalize_whitespace(data.decode("utf-8", errors="ignore"))
    refs = [EvidenceRef(file_name=attachment.name, locator="raw", snippet=text[:220] or attachment.name)]
    return text, refs


def _classify_document(attachment: Attachment, text: str, rules: dict[str, Any]) -> str:
    lowered_name = attachment.name.lower()
    lowered_text = (text or "").lower()
    best_type = "Unknown"
    best_score = 0
    for document_type, config in rules.get("document_types", {}).items():
        score = 0
        for pattern in config.get("filename_patterns", []):
            if pattern.lower() in lowered_name:
                score += 3
        for pattern in config.get("text_patterns", []):
            if pattern.lower() in lowered_text:
                score += 2
        if score > best_score:
            best_score = score
            best_type = document_type
    return best_type


def _normalize_date(value: str | None) -> str | None:
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%B %d, %Y", "%b %d, %Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(value.strip(), fmt).date().isoformat()
        except ValueError:
            continue
    return None


def _extract_fields(document_type: str, text: str, rules: dict[str, Any]) -> dict[str, Any]:
    extracted: dict[str, Any] = {}
    extractor_rules = rules.get("field_extractors", {})
    for field_name, config in extractor_rules.items():
        applies_to = config.get("document_types")
        if applies_to and document_type not in applies_to:
            continue
        for pattern in config.get("regexes", []):
            match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
            if not match:
                continue
            value = match.group(1).strip() if match.groups() else match.group(0).strip()
            if field_name.endswith("_date"):
                value = _normalize_date(value) or value
            extracted[field_name] = value
            break
    return extracted


def _group_field_values(documents: list[ReviewedDocument], field_name: str, document_types: list[str]) -> list[tuple[str, Any]]:
    values: list[tuple[str, Any]] = []
    for document in documents:
        if document.document_type not in document_types:
            continue
        value = document.fields.get(field_name)
        if value:
            values.append((document.file_name, value))
    return values


def _cross_check_details(values: list[tuple[str, Any]]) -> str:
    return "; ".join(f"{file_name}={value}" for file_name, value in values)


def _shareholding_consistent(documents: list[ReviewedDocument]) -> bool:
    shareholder_names = _group_field_values(documents, "shareholder_name", ["COI", "OnboardingForm"])
    if len(shareholder_names) < 2:
        return False
    normalized_names = {_normalize_entity(value) for _, value in shareholder_names}
    if len(normalized_names) != 1:
        return False
    holdings = _group_field_values(documents, "shareholding_percent", ["COI", "OnboardingForm"])
    if len(holdings) < 2:
        return False
    contains_percent = any("%" in str(value) for _, value in holdings)
    contains_shares = any("share" in str(value).lower() for _, value in holdings)
    return contains_percent and contains_shares


def _run_cross_checks(documents: list[ReviewedDocument], rules: dict[str, Any], reference_date: datetime) -> tuple[list[CrossCheckResult], list[str]]:
    checks: list[CrossCheckResult] = []
    rule_hits: list[str] = []
    for required in rules.get("required_documents", []):
        matched_documents = [document.file_name for document in documents if document.document_type == required]
        present = bool(matched_documents)
        details = f"required document {required} found in {', '.join(matched_documents)}" if present else f"required document {required} is missing"
        checks.append(CrossCheckResult(check=f"required_{required}", result="pass" if present else "fail", details=details))
        if present:
            rule_hits.append(f"required:{required}")

    for comparison in rules.get("comparisons", []):
        field_name = comparison["field"]
        values = _group_field_values(documents, field_name, comparison["document_types"])
        if len(values) < 2:
            checks.append(CrossCheckResult(check=comparison["id"], result="unknown", details="not enough evidence"))
            continue

        if comparison["id"] == "shareholding_consistency" and _shareholding_consistent(documents):
            checks.append(
                CrossCheckResult(
                    check=comparison["id"],
                    result="pass",
                    details=f"sole-shareholder pattern matched across {', '.join(file_name for file_name, _ in values)}",
                )
            )
            rule_hits.append(comparison["id"])
            continue

        normalized_values = {_normalize_entity(value) for _, value in values}
        if len(normalized_values) == 1:
            checks.append(CrossCheckResult(check=comparison["id"], result="pass", details=f"matched across {_cross_check_details(values)}"))
            rule_hits.append(comparison["id"])
        else:
            checks.append(CrossCheckResult(check=comparison["id"], result="fail", details=_cross_check_details(values)))

    freshness_rules = rules.get("freshness_rules", {})
    for document in documents:
        max_age = freshness_rules.get(document.document_type)
        issue_date = document.fields.get("issue_date")
        if not max_age or not issue_date:
            continue
        normalized = _normalize_date(str(issue_date))
        if normalized is None:
            checks.append(
                CrossCheckResult(
                    check=f"freshness_{document.document_type}",
                    result="unknown",
                    details=f"{document.file_name}: issue_date is not parseable",
                )
            )
            continue
        age_days = (reference_date.date() - datetime.fromisoformat(normalized).date()).days
        if age_days <= int(max_age):
            checks.append(
                CrossCheckResult(
                    check=f"freshness_{document.document_type}",
                    result="pass",
                    details=f"{document.file_name}: {age_days} days old",
                )
            )
            rule_hits.append(f"freshness:{document.document_type}")
        else:
            checks.append(
                CrossCheckResult(
                    check=f"freshness_{document.document_type}",
                    result="fail",
                    details=f"{document.file_name}: {age_days} days old exceeds {max_age}",
                )
            )
    return checks, rule_hits


def _recommend(documents: list[ReviewedDocument], checks: list[CrossCheckResult]) -> ReviewRecommendation:
    failed = [check for check in checks if check.result == "fail"]
    if any(check.check in {"identity_consistency", "passport_consistency"} for check in failed):
        return ReviewRecommendation(decision="reject", confidence=0.92, reasons=[check.details for check in failed])
    if any(check.check.startswith("required_") for check in failed) or any(check.check.startswith("freshness_") for check in failed):
        return ReviewRecommendation(decision="resubmit", confidence=0.78, reasons=[check.details for check in failed])
    if documents and not failed and len([check for check in checks if check.result == "pass"]) >= 3:
        return ReviewRecommendation(decision="approve", confidence=0.81, reasons=["core checks passed across uploaded documents"])
    return ReviewRecommendation(decision="manual_review", confidence=0.55, reasons=[check.details for check in failed] or ["additional human review required"])


def review_uploaded_documents(files: list[Attachment], package: PlatformPackage, reference_date: datetime) -> DocumentReviewResult:
    rules = package.rules.get("document_rules", {})
    documents: list[ReviewedDocument] = []
    for attachment in files:
        data = _read_attachment_bytes(attachment)
        text, evidence_refs = _extract_text(attachment, data)
        document_type = _classify_document(attachment, text, rules)
        fields = _extract_fields(document_type, text, rules)
        documents.append(
            ReviewedDocument(
                file_name=attachment.name,
                document_type=document_type,
                fields=fields,
                evidence_refs=_attach_file_name(attachment.name, evidence_refs),
            )
        )
    checks, rule_hits = _run_cross_checks(documents, rules, reference_date)
    recommendation = _recommend(documents, checks)
    return DocumentReviewResult(documents=documents, cross_checks=checks, recommendation=recommendation, rule_hits=rule_hits)
